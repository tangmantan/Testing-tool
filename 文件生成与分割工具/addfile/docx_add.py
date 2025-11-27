import os
from docx import Document

def generate_fixed_size_docx(filename, target_size_mb):
    """
    快速生成指定大小的 docx 文件 (通过二进制填充)
    :param filename: 输出文件名
    :param target_size_mb: 目标大小 (MB)
    """
    
    # 1. 先生成一个合法的、最小的基础 docx 文件
    doc = Document()
    doc.add_heading('文件大小测试', 0)
    doc.add_paragraph(f'这是一个自动生成的测试文件。目标大小: {target_size_mb} MB。')
    doc.save(filename)

    # 2. 计算需要填充的字节数
    current_size = os.path.getsize(filename)
    target_bytes = int(target_size_mb * 1024 * 1024)
    padding_size = target_bytes - current_size

    if padding_size <= 0:
        print(f"⚠️ 警告: 初始文件 ({current_size/1024/1024:.2f} MB) 已超过目标大小。")
        return

    print(f"正在填充数据以达到 {target_size_mb} MB...")

    # 3. 以二进制追加模式 ('ab') 打开文件并填充空字节
    # 使用分块写入，防止生成大文件时内存溢出
    chunk_size = 1024 * 1024 * 10 # 每次写 10MB
    with open(filename, 'ab') as f:
        while padding_size > 0:
            write_size = min(padding_size, chunk_size)
            # 写入空字节 (b'\0')
            f.write(b'\0' * write_size)
            padding_size -= write_size

    final_size = os.path.getsize(filename) / (1024 * 1024)
    print(f"✅ 成功! 文件: {filename}, 最终大小: {final_size:.2f} MB")

if __name__ == "__main__":
    # 在这里修改文件名和大小
    generate_fixed_size_docx("测试文档_5MB.docx", 5)   # 生成 10MB