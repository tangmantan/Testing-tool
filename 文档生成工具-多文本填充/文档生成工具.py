# -*- coding: utf-8 -*-
"""
-------------------------------------------------
   File Name：     文档生成工具
   Description:    生成指定大小的不同格式文档
   Author :       zh
   date：          2025/11/14
-------------------------------------------------
"""
import os
import random
import string
import sys
import ctypes
import tempfile
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfMerger


app = Flask(__name__)

# 确保输出目录存在
OUTPUT_DIR = 'generated_files'
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# 随机文本库
SAMPLE_TEXTS = [
    "燕子去了，有再来的时候；杨柳枯了，有再青的时候；桃花谢了，有再开的时候。",
    "春眠不觉晓，处处闻啼鸟。夜来风雨声，花落知多少。",
    "床前明月光，疑是地上霜。举头望明月，低头思故乡。",
    "白日依山尽，黄河入海流。欲穷千里目，更上一层楼。",
    "慈母手中线，游子身上衣。临行密密缝，意恐迟迟归。",
    "红豆生南国，春来发几枝。愿君多采撷，此物最相思。",
    "好雨知时节，当春乃发生。随风潜入夜，润物细无声。",
    "千山鸟飞绝，万径人踪灭。孤舟蓑笠翁，独钓寒江雪。",
    "空山新雨后，天气晚来秋。明月松间照，清泉石上流。",
    "君自故乡来，应知故乡事。来日绮窗前，寒梅著花未。"
]

def generate_random_text(target_size_bytes):
    """生成随机文本内容"""
    text = ""
    current_size = 0
    
    while current_size < target_size_bytes:
        # 随机选择一段文本
        text += random.choice(SAMPLE_TEXTS)
        text += "\n"
        
        # 随机添加一些数字
        if random.random() > 0.7:
            text += ''.join(random.choices(string.digits, k=random.randint(10, 50)))
            text += "\n"
        
        current_size = len(text.encode('utf-8'))
    
    return text[:target_size_bytes]

def generate_txt_file(filename, size_mb):
    filepath = os.path.join(OUTPUT_DIR, filename)
    target_size = int(size_mb * 1024 * 1024)
    print(f"目标大小: {size_mb}MB")

    lines = []  # 用列表存储每一行，避免频繁字符串拼接
    current_size = 0
    iteration = 0
    check_interval = 1000  # 每1000次迭代打印一次进度，可调整

    while current_size < target_size:
        iteration += 1

        # 添加随机文本行
        line = random.choice(SAMPLE_TEXTS) + "\n"
        lines.append(line)
        current_size += len(line.encode('utf-8'))

        # 随机添加数字行
        if random.random() > 0.7:
            num_line = ''.join(random.choices(string.digits, k=random.randint(10, 50))) + "\n"
            lines.append(num_line)
            current_size += len(num_line.encode('utf-8'))

        # 打印进度
        if iteration % check_interval == 0 or current_size >= target_size:
            progress = (current_size / target_size * 100) if target_size > 0 else 0
            print(f"进度: {current_size / 1024 / 1024:.2f}MB / {size_mb}MB ({progress:.1f}%, 迭代{iteration}次)")

    # 最终确保不超过目标大小
    final_text = ''.join(lines)
    if len(final_text.encode('utf-8')) > target_size:
        final_text = final_text.encode('utf-8')[:target_size].decode('utf-8', errors='ignore')

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(final_text)

    actual_size = os.path.getsize(filepath)
    print(f"生成完成: {actual_size / 1024 / 1024:.2f}MB, 共{iteration}次迭代")
    return filepath, actual_size
    


def generate_docx_file(filename, size_mb):
    """生成DOCX文件"""
    filepath = os.path.join(OUTPUT_DIR, filename)
    target_size = int(size_mb * 1024 * 1024)
    
    doc = Document()
    doc.add_heading('随机生成文档', 0)
    
    current_size = 0
    iteration = 0

    print("如果要精确控制文件大小，需要改小randint的范围，但相应的速度会更慢")
    # 循环添加内容直到达到目标大小
    while current_size < target_size:
        iteration += 1
        
     
        
        # 添加段落 - 大幅增加每段的内容量，克服DOCX压缩
        text = random.choice(SAMPLE_TEXTS) * random.randint(5000, 10000)  
        doc.add_paragraph(text)
        
        
        # 动态调整检查频率：开始时每500段检查，接近目标时每10段检查
        check_interval = 100 if current_size > target_size * 0.99 else 500
        
        if iteration % check_interval == 0:
            doc.save(filepath)
            current_size = os.path.getsize(filepath)
            progress = (current_size / target_size * 100) if target_size > 0 else 0
            print(f"DOCX生成进度: {current_size / 1024 / 1024:.2f}MB / {size_mb}MB ({progress:.1f}%, 迭代{iteration}次)")
            
            # 如果已经超过目标大小，退出
            if current_size >= target_size:
                break
    
    # 最终保存
    doc.save(filepath)
    actual_size = os.path.getsize(filepath)
    print(f"DOCX生成完成: {actual_size / 1024 / 1024:.2f}MB (共{iteration}次迭代)")
    

    return filepath, actual_size







def generate_pdf_file(filename, size_mb):
    """生成PDF文件 - 支持中文字体"""
    filepath = os.path.join(OUTPUT_DIR, filename)
    target_size = int(size_mb * 1024 * 1024)
    
    # 确保文件不存在
    if os.path.exists(filepath):
        os.remove(filepath)
    
    # 尝试注册中文字体
    try:
        # 尝试使用系统中常见的中文字体
        font_paths = [
            "C:/Windows/Fonts/simhei.ttf",      # 黑体
            "C:/Windows/Fonts/simsun.ttc",      # 宋体
            "C:/Windows/Fonts/msyh.ttc",        # 微软雅黑
            "C:/Windows/Fonts/simkai.ttf",      # 楷体
        ]
        
        font_registered = False
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    # 从TTC文件中提取第一个字体
                    if font_path.endswith('.ttc'):
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                    else:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    font_registered = True
                    print(f"成功注册中文字体: {font_path}")
                    break
                except Exception as e:
                    print(f"注册字体失败 {font_path}: {e}")
                    continue
        
        if not font_registered:
            print("警告: 无法注册中文字体，将使用英文字体")
            use_chinese = False
        else:
            use_chinese = True
    except Exception as e:
        print(f"字体注册出错: {e}")
        use_chinese = False
    
    # 英文示例文本
    ENGLISH_SAMPLE_TEXTS = [
        "The quick brown fox jumps over the lazy dog. This pangram sentence contains every letter of the alphabet at least once.",
        "To be or not to be, that is the question. Whether 'tis nobler in the mind to suffer the slings and arrows of outrageous fortune.",
        "It was the best of times, it was the worst of times, it was the age of wisdom, it was the age of foolishness.",
        "All happy families are alike; each unhappy family is unhappy in its own way.",
        "Call me Ishmael. Some years ago—never mind how long precisely—having little or no money in my purse.",
        "In the beginning God created the heavens and the earth. Now the earth was formless and empty.",
        "It is a truth universally acknowledged, that a single man in possession of a good fortune, must be in want of a wife.",
        "Many years later, as he faced the firing squad, Colonel Aureliano Buendía was to remember that distant afternoon.",
        "Happy families are all alike; every unhappy family is unhappy in its own way.",
        "You don't know about me without you have read a book by the name of The Adventures of Tom Sawyer."
    ]
    
    print(f"目标大小: {size_mb}MB")
    
    # 使用临时文件策略
    temp_files = []
    total_pages = 0
    actual_size = 0
    
    # 循环生成临时PDF文件，直到达到目标大小
    while actual_size < target_size:
        # 创建临时文件
        temp_file = os.path.join(tempfile.gettempdir(), f"temp_pdf_{random.randint(1000, 9999)}.pdf")
        temp_files.append(temp_file)
        
        # 创建临时PDF
        c = canvas.Canvas(temp_file, pagesize=A4)
        width, height = A4
        
        # 生成10页内容
        for page_num in range(10):
            total_pages += 1
            
            # 不添加页面标题，直接从页面顶部开始内容
            y_position = height - 0.3 * inch  # 更接近页面顶部
            
            # 每页添加适量的文本
            for line_num in range(200):  # 200行
                # 根据是否支持中文选择文本
                if use_chinese:
                    # 使用中文文本
                    text = random.choice(SAMPLE_TEXTS)  # 不重复，避免过长
                    text += "中文文本测试内容，中文文本测试内容。"

                    # 设置中文字体
                    c.setFont('ChineseFont', 8)
                else:
                    # 使用英文文本
                    text = random.choice(ENGLISH_SAMPLE_TEXTS)  # 不重复，避免过长
                    text += ''.join(random.choices(string.ascii_letters + string.digits, k=20))
                    # 使用默认字体
                    c.setFont('Helvetica', 8)
                
                # 确保文本不会超出页面边界 - 添加文本换行处理
                if use_chinese:
                    text_width = c.stringWidth(text, 'ChineseFont', 8)
                else:
                    text_width = c.stringWidth(text, 'Helvetica', 8)
                
                # 如果文本太宽，进行换行处理
                max_width = width - 0.4 * inch
                if text_width > max_width:
                    # 计算每行最大字符数
                    if use_chinese:
                        # 中文字符宽度大致相等，按比例计算
                        max_chars = int(len(text) * max_width / text_width)
                        if max_chars < 10:  # 确保至少有10个字符
                            max_chars = 10
                    else:
                        # 英文字符宽度不等，保守估计
                        max_chars = int(len(text) * max_width / text_width * 0.9)  # 乘以0.9作为安全系数
                        if max_chars < 15:  # 确保至少有15个字符
                            max_chars = 15
                    
                    # 截断文本并添加省略号
                    if len(text) > max_chars:
                        text = text[:max_chars-3] + "..."
                
                # 确保Y位置不会超出页面底部
                if y_position < 0.5 * inch:  # 保留底部空间
                    break
                
                c.drawString(0.2 * inch, y_position, text)
                y_position -= 0.15 * inch  # 增加行间距，确保不重叠
                
                if y_position < 0.5 * inch:  # 保留底部空间
                    break
            
            # 在添加完内容后再添加新页面
            c.showPage()
        
        # 保存临时文件
        c.save()
        
        # 检查当前总大小
        actual_size = sum(os.path.getsize(f) for f in temp_files if os.path.exists(f))
        print(f"PDF生成进度: 已生成 {total_pages} 页, 当前大小: {actual_size / 1024 / 1024:.2f}MB")
        
        # 如果接近目标大小，减少生成页数
        if actual_size >= target_size * 0.98:
            print("接近目标大小，开始精确控制...")
            break
    
    # 合并所有临时文件
    merger = PdfMerger()
    for temp_file in temp_files:
        if os.path.exists(temp_file):
            merger.append(temp_file)
    
    # 保存最终文件
    merger.write(filepath)
    merger.close()
    
    # 清理临时文件
    for temp_file in temp_files:
        if os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except:
                pass
    
    # 检查最终文件大小
    actual_size = os.path.getsize(filepath)
    print(f"PDF最终大小: {actual_size / 1024 / 1024:.2f}MB, 总页数: {total_pages}")
    
    return filepath, actual_size










@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_file():
    """生成文件"""
    try:
        file_format = request.form.get('format', 'txt')
        size_mb = float(request.form.get('size', 1))
        
        # 限制文件大小
        if size_mb > 100:
            return jsonify({'error': '文件大小不能超过100MB'}), 400
        
        if size_mb <= 0:
            return jsonify({'error': '文件大小必须大于0'}), 400
        
        # 生成文件名
        timestamp = random.randint(1000, 9999)
        filename = f'random_file_{timestamp}.{file_format}'
        
        # 根据格式生成文件
        if file_format == 'txt':
            filepath, actual_size = generate_txt_file(filename, size_mb)
        elif file_format == 'docx':
            filepath, actual_size = generate_docx_file(filename, size_mb)
        elif file_format == 'pdf':
            filepath, actual_size = generate_pdf_file(filename, size_mb)
        else:
            return jsonify({'error': '不支持的文件格式'}), 400
        
        # 返回下载链接
        return jsonify({
            'success': True,
            'filename': filename,
            'actual_size': f'{actual_size / 1024 / 1024:.2f} MB',
            'download_url': f'/download/{filename}'
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """下载文件"""
    filepath = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True, download_name=filename)
    else:
        return "文件不存在", 404

@app.route('/list')
def list_files():
    """列出所有生成的文件"""
    files = []
    for filename in os.listdir(OUTPUT_DIR):
        filepath = os.path.join(OUTPUT_DIR, filename)
        size = os.path.getsize(filepath)
        files.append({
            'name': filename,
            'size': f'{size / 1024 / 1024:.2f} MB',
            'download_url': f'/download/{filename}',
            'delete_url': f'/delete/{filename}'
        })
    return jsonify(files)

@app.route('/delete/<filename>', methods=['DELETE'])
def delete_file(filename):
    """删除文件"""
    filepath = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(filepath):
        try:
            os.remove(filepath)
            return jsonify({'success': True, 'message': f'文件 {filename} 已删除'})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    else:
        return jsonify({'success': False, 'error': '文件不存在'}), 404

if __name__ == '__main__':
    import sys
    
    # 解决Windows CMD下的Quick Edit Mode问题
    # if sys.platform == 'win32':
    #     import ctypes
    #     kernel32 = ctypes.windll.kernel32
    #     kernel32.SetConsoleMode(kernel32.GetStdHandle(-10), 128)
    
    print("=" * 50)
    print("文档生成工具已启动！")
    print("请在浏览器中打开: http://127.0.0.1:5000")
    print("按 Ctrl+C 停止服务")
    print("提示: 如果页面无响应，请不要在CMD窗口中选择文字")
    print("=" * 50)
    
    # 使用threaded=True来提高响应性
    app.run(debug=True, host='0.0.0.0', port=5000, threaded=True)
